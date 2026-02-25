from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
import shutil
import os
import subprocess
import uuid
from pathlib import Path
import re
import tempfile
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageChops, ImageDraw, ImageFont

router = APIRouter(prefix="/md-to-docx", tags=["Markdown to DOCX"])

UPLOAD_DIR = Path("./tmp/uploads")
OUTPUT_DIR = Path("./tmp/outputs")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _set_cell_borders(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    for edge in ('top', 'left', 'bottom', 'right'):
        element = tc_borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tc_borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '8')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')


def _set_cell_shading(cell, fill='D9E2F3'):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)


def _set_header_row_style(table):
    if not table.rows:
        return

    for cell in table.rows[0].cells:
        _set_cell_shading(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)


def _enforce_document_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)

    heading_1 = doc.styles['Heading 1']
    heading_1.font.name = 'Arial'
    heading_1.font.size = Pt(24)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(0, 51, 102)

    heading_2 = doc.styles['Heading 2']
    heading_2.font.name = 'Arial'
    heading_2.font.size = Pt(18)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(0, 85, 164)

    heading_3 = doc.styles['Heading 3']
    heading_3.font.name = 'Arial'
    heading_3.font.size = Pt(14)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(0, 51, 102)


def _trim_image_file(image_path: Path):
    try:
        with Image.open(image_path) as img:
            original_width, original_height = img.size

            if img.mode != 'RGB':
                rgb = img.convert('RGB')
            else:
                rgb = img

            white_bg = Image.new('RGB', rgb.size, 'white')
            diff = ImageChops.difference(rgb, white_bg)
            bbox = diff.getbbox()

            if bbox is None:
                return

            left, top, right, bottom = bbox
            padding = 4
            left = max(0, left - padding)
            top = max(0, top - padding)
            right = min(original_width, right + padding)
            bottom = min(original_height, bottom + padding)

            if (right - left) >= original_width and (bottom - top) >= original_height:
                return

            width_ratio = (right - left) / original_width
            height_ratio = (bottom - top) / original_height

            if width_ratio > 0.98 and height_ratio > 0.98:
                return

            cropped = img.crop((left, top, right, bottom))
            cropped.save(image_path)
    except Exception as error:
        print(f"Image trim skipped for {image_path.name}: {error}")


def _trim_docx_media_images(docx_path: Path):
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        with zipfile.ZipFile(docx_path, 'r') as source_zip:
            source_zip.extractall(temp_root)

        media_dir = temp_root / 'word' / 'media'
        if media_dir.exists():
            for image_file in media_dir.iterdir():
                if image_file.suffix.lower() in ('.png', '.jpg', '.jpeg'):
                    _trim_image_file(image_file)

        with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as target_zip:
            for file_path in temp_root.rglob('*'):
                if file_path.is_file():
                    target_zip.write(file_path, file_path.relative_to(temp_root))


def _sync_inline_shape_aspect_ratio(docx_path: Path):
    doc = Document(str(docx_path))
    max_width = None
    if doc.sections:
        widths = []
        for section in doc.sections:
            widths.append(int(section.page_width - section.left_margin - section.right_margin))
        if widths:
            max_width = min(widths)

    settings = doc.settings._element
    no_compress = settings.find(qn('w:doNotCompressPictures'))
    if no_compress is None:
        no_compress = OxmlElement('w:doNotCompressPictures')
        settings.append(no_compress)
    no_compress.set(qn('w:val'), 'true')

    for shape in doc.inline_shapes:
        try:
            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rel_id = blip.embed
            image_part = doc.part.related_parts[rel_id]
            pixel_width = image_part.image.px_width
            pixel_height = image_part.image.px_height
            if pixel_width and pixel_height:
                target_width = int(shape.width)
                if max_width is not None:
                    target_width = max_width

                shape.width = Emu(target_width)
                corrected_height = int(target_width * pixel_height / pixel_width)
                shape.height = Emu(corrected_height)
        except Exception as error:
            print(f"Inline shape ratio sync skipped: {error}")
    doc.save(str(docx_path))


def _best_diagram_layout(img_width_px, img_height_px, avail_portrait, avail_landscape):
    candidates = [
        ('portrait', False, avail_portrait),
        ('portrait', True, avail_portrait),
        ('landscape', False, avail_landscape),
        ('landscape', True, avail_landscape),
    ]

    best = None
    for orientation, rotate_90, (avail_w, avail_h) in candidates:
        render_w = img_height_px if rotate_90 else img_width_px
        render_h = img_width_px if rotate_90 else img_height_px

        if render_w <= 0 or render_h <= 0:
            continue

        scale = min(avail_w / render_w, avail_h / render_h)
        display_w = int(render_w * scale)
        display_h = int(render_h * scale)
        area = display_w * display_h

        option = {
            'orientation': orientation,
            'rotate_90': rotate_90,
            'width_emu': display_w,
            'height_emu': display_h,
            'area': area,
        }

        if best is None or option['area'] > best['area']:
            best = option

    return best


def _is_large_diagram(img_width_px, img_height_px):
    shorter_edge = min(img_width_px, img_height_px)
    return (img_width_px * img_height_px) >= 250000 and shorter_edge >= 220


def _create_internal_hyperlink(anchor_name, link_text):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_name)
    hyperlink.set(qn('w:history'), '1')

    run = OxmlElement('w:r')
    run_props = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    run_props.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    run_props.append(underline)
    run.append(run_props)

    text_el = OxmlElement('w:t')
    text_el.text = link_text
    run.append(text_el)
    hyperlink.append(run)
    return hyperlink


def _insert_reference_after_paragraph(paragraph, anchor_name, figure_label):
    ref_paragraph = OxmlElement('w:p')

    prefix_run = OxmlElement('w:r')
    prefix_text = OxmlElement('w:t')
    prefix_text.text = 'See full-page '
    prefix_run.append(prefix_text)
    ref_paragraph.append(prefix_run)

    ref_paragraph.append(_create_internal_hyperlink(anchor_name, figure_label))

    suffix_run = OxmlElement('w:r')
    suffix_text = OxmlElement('w:t')
    suffix_text.text = ' in Appendix A.'
    suffix_run.append(suffix_text)
    ref_paragraph.append(suffix_run)

    paragraph._p.addnext(ref_paragraph)


def _add_bookmark_to_paragraph(paragraph, bookmark_name, bookmark_id):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _inject_appendix_title_into_image(image_path: Path, title_text: str):
    try:
        with Image.open(image_path) as source_image:
            image = source_image.convert('RGB')
            width, height = image.size

            banner_height = max(80, int(width * 0.06))
            canvas = Image.new('RGB', (width, height + banner_height), 'white')
            draw = ImageDraw.Draw(canvas)

            draw.rectangle((0, 0, width, banner_height), fill=(242, 246, 255))

            try:
                font = ImageFont.truetype('arial.ttf', max(24, int(width * 0.03)))
            except Exception:
                font = ImageFont.load_default()

            text_bbox = draw.textbbox((0, 0), title_text, font=font)
            text_w = text_bbox[2] - text_bbox[0]
            text_h = text_bbox[3] - text_bbox[1]
            text_x = max(12, (width - text_w) // 2)
            text_y = max(8, (banner_height - text_h) // 2)
            draw.text((text_x, text_y), title_text, fill=(0, 51, 102), font=font)

            canvas.paste(image, (0, banner_height))
            canvas.save(image_path)
    except Exception as error:
        print(f"Appendix title injection skipped for {image_path.name}: {error}")


def _append_full_page_diagram_appendix(docx_path: Path):
    doc = Document(str(docx_path))
    if not doc.inline_shapes:
        doc.save(str(docx_path))
        return

    first_section = doc.sections[0]
    base_page_w = int(first_section.page_width)
    base_page_h = int(first_section.page_height)
    base_left = int(first_section.left_margin)
    base_right = int(first_section.right_margin)
    base_top = int(first_section.top_margin)
    base_bottom = int(first_section.bottom_margin)

    portrait_page_w = min(base_page_w, base_page_h)
    portrait_page_h = max(base_page_w, base_page_h)
    landscape_page_w = portrait_page_h
    landscape_page_h = portrait_page_w

    avail_portrait = (
        portrait_page_w - base_left - base_right,
        portrait_page_h - base_top - base_bottom,
    )
    avail_landscape = (
        landscape_page_w - base_left - base_right,
        landscape_page_h - base_top - base_bottom,
    )

    diagram_entries = []
    seen_rel_ids = set()
    for shape in list(doc.inline_shapes):
        try:
            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rel_id = blip.embed
            if rel_id in seen_rel_ids:
                continue
            seen_rel_ids.add(rel_id)
            image_part = doc.part.related_parts[rel_id]
            img_w = image_part.image.px_width
            img_h = image_part.image.px_height
            if not _is_large_diagram(img_w, img_h):
                continue
            diagram_entries.append({
                'rel_id': rel_id,
                'image_part': image_part,
                'img_w': img_w,
                'img_h': img_h,
                'shape': shape,
            })
        except Exception as error:
            print(f"Appendix collection skipped a shape: {error}")

    if not diagram_entries:
        doc.save(str(docx_path))
        return

    figure_map = {}
    for index, entry in enumerate(diagram_entries, start=1):
        figure_map[entry['rel_id']] = {
            'label': f'Appendix Figure A{index}',
            'anchor': f'appendix_figure_a{index}',
            'index': index,
        }

    annotated_rel_ids = set()
    for paragraph in doc.paragraphs:
        current_rel_id = None
        for run in paragraph.runs:
            for blip in run._r.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                rel_id = blip.get(qn('r:embed'))
                if rel_id in figure_map:
                    current_rel_id = rel_id
                    break
            if current_rel_id:
                break

        if current_rel_id and current_rel_id not in annotated_rel_ids:
            meta = figure_map[current_rel_id]
            _insert_reference_after_paragraph(paragraph, meta['anchor'], meta['label'])
            annotated_rel_ids.add(current_rel_id)

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)

        for entry in diagram_entries:
            try:
                index = figure_map[entry['rel_id']]['index']
                image_part = entry['image_part']
                source_ext = image_part.filename.split('.')[-1].lower()
                if source_ext not in ('png', 'jpg', 'jpeg'):
                    continue

                source_path = temp_root / f'diagram_{index}.{source_ext}'
                source_path.write_bytes(image_part.blob)

                img_w = entry['img_w']
                img_h = entry['img_h']

                layout = _best_diagram_layout(img_w, img_h, avail_portrait, avail_landscape)
                if layout is None:
                    continue

                render_path = source_path
                if layout['rotate_90']:
                    with Image.open(source_path) as source_image:
                        rotated = source_image.rotate(90, expand=True)
                        render_path = temp_root / f'diagram_{index}_rotated.{source_ext}'
                        rotated.save(render_path)

                figure_label = figure_map[entry['rel_id']]['label']
                _inject_appendix_title_into_image(render_path, figure_label)
                with Image.open(render_path) as titled_image:
                    titled_w, titled_h = titled_image.size

                layout = _best_diagram_layout(titled_w, titled_h, avail_portrait, avail_landscape)
                if layout is None:
                    continue

                section = doc.add_section(WD_SECTION_START.NEW_PAGE)
                section.left_margin = Emu(base_left)
                section.right_margin = Emu(base_right)
                section.top_margin = Emu(base_top)
                section.bottom_margin = Emu(base_bottom)

                if layout['orientation'] == 'landscape':
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width = Emu(landscape_page_w)
                    section.page_height = Emu(landscape_page_h)
                else:
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width = Emu(portrait_page_w)
                    section.page_height = Emu(portrait_page_h)

                picture_paragraph = doc.add_paragraph()
                _add_bookmark_to_paragraph(
                    picture_paragraph,
                    figure_map[entry['rel_id']]['anchor'],
                    9000 + index
                )
                run = picture_paragraph.add_run()
                run.add_picture(
                    str(render_path),
                    width=Emu(layout['width_emu']),
                    height=Emu(layout['height_emu'])
                )

            except Exception as error:
                print(f"Appendix render skipped diagram {index}: {error}")

    doc.save(str(docx_path))


def _apply_table_style(table):
    preferred_styles = ['MyCustomTable', 'Table Grid', 'Normal Table']
    for style_name in preferred_styles:
        try:
            table.style = style_name
            return
        except Exception:
            continue


def _enforce_table_borders(docx_path: Path):
    doc = Document(str(docx_path))
    _enforce_document_styles(doc)
    for table in doc.tables:
        _apply_table_style(table)
        _set_header_row_style(table)
        for row in table.rows:
            for cell in row.cells:
                _set_cell_borders(cell)
    doc.save(str(docx_path))
    _trim_docx_media_images(docx_path)
    _sync_inline_shape_aspect_ratio(docx_path)
    _append_full_page_diagram_appendix(docx_path)

# Helper: Preprocess Markdown for Mermaid
def preprocess_markdown(file_path):
    """
    Reads markdown file, converts <div class="mermaid"> to fenced code blocks,
    and returns path to a temporary file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Regex to transform <div class="mermaid">...</div> into ```mermaid...```
        pattern = re.compile(r'<div class="mermaid">\s*(.*?)\s*</div>', re.DOTALL)
        
        def replacement(match):
            code = match.group(1).strip()
            return f"\n```mermaid\n{code}\n```\n"
        
        new_content = pattern.sub(replacement, content)
        
        base = os.path.splitext(file_path)[0]
        temp_path = f"{base}_processed.md"
        
        with open(temp_path, 'w', encoding='utf-8') as f:
            f.write(new_content)
        
        return temp_path
    except Exception as e:
        print(f"Error preprocessing markdown: {e}")
        return file_path

@router.post("/convert/")
async def convert_markdown_to_docx(file: UploadFile = File(...)):
    if not file.filename.endswith(".md"):
        raise HTTPException(status_code=400, detail="Only .md files are allowed")

    request_id = str(uuid.uuid4())
    input_path = UPLOAD_DIR / f"{request_id}_{file.filename}"
    output_filename = f"{Path(file.filename).stem}.docx"
    output_path = OUTPUT_DIR / f"{request_id}_{output_filename}"

    # Save uploaded file
    with open(input_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Preprocess
    processed_path = preprocess_markdown(str(input_path))

    # Puppeteer Config for Mermaid Filter
    # We want a high density (scale factor) but we don't want a huge fixed viewport 
    # that forces whitespace if the diagram is small.
    puppeteer_config = {
        "executablePath": "/usr/bin/google-chrome",
        "args": ["--no-sandbox", "--disable-setuid-sandbox"],
        # Render denser images to preserve sharpness after any post-processing.
        "defaultViewport": {"width": 2200, "height": 1400, "deviceScaleFactor": 8}
    }
    import json
    # mermaid-filter expects .puppeteer.json in CWD
    config_path = Path(".puppeteer.json")
    with open(config_path, "w") as f:
        json.dump(puppeteer_config, f)

    # Mermaid Config to improve quality (increase scale)
    # mermaid-filter looks for .mermaid-config.json in CWD
    mermaid_config = {
        "theme": "default", 
        "startOnLoad": False,  
        "themeVariables": {
            "fontFamily": "Arial",
            "fontSize": "22px"
        },
        "flowchart": {
            "useMaxWidth": True, # Prevents massive SVGs that might break things, lets Mermaid manage width
            "htmlLabels": True,
            "curve": "cardinal"
        },
        "sequence": {
             "useMaxWidth": True
        }
    }
    # Also we can set puppeterr options often in the same place or separate
    # But usually mermaid-filter has a separate mechanism for scale.
    # It seems mermaid-filter 1.4.x might determine scale via puppeteer viewport or explicit args?
    # Actually, recent versions respecting --width or similar in the code block, 
    # but global settings are in .mermaid-filter.json? No, it's .puppeteer.json usually acting as the bridge.
    # Let's try creating a .mermaid-config.json
    mermaid_conf_path = Path(".mermaid-config.json")
    with open(mermaid_conf_path, "w") as f:
        json.dump(mermaid_config, f)

    # Convert using Pandoc
    # Command: pandoc input.md -o output.docx -F mermaid-filter --reference-doc=reference.docx (if exists)
    # We add --verbose to see mermaid-filter logs
    # filter_table_style.lua forces 'TableGrid' style on all tables
    filter_style_path = Path("md-to-docx/filter_table_style.lua")
    if not filter_style_path.exists():
        filter_style_path = Path("filter_table_style.lua") # if in cwd

    cmd = [
        "pandoc",
        "-f", "gfm+hard_line_breaks+raw_html",
        processed_path,
        "-o", str(output_path),
        "-F", "mermaid-filter",
        "--verbose"
    ]
    
    if filter_style_path.exists():
         cmd.extend(["--lua-filter", str(filter_style_path)])

    # Remove MERMAID_FILTER env vars that interfere with manual config
    # We rely on the config files above.
    env = os.environ.copy()
    env["MERMAID_FILTER_SCALE"] = "4"
    # env["MERMAID_FILTER_WIDTH"] = "900" # Let it flow naturally
    
    # Check for reference doc in current dir (where app is running) or md-to-docx folder
    ref_doc_name = "reference.docx"
    ref_doc_path = Path(ref_doc_name)
    if not ref_doc_path.exists():
         # Check inside md-to-docx folder relative to current working directory
         possible_path = Path("md-to-docx") / ref_doc_name
         if possible_path.exists():
             ref_doc_path = possible_path
    
    if ref_doc_path.exists():
        cmd.extend(["--reference-doc", str(ref_doc_path)])

    try:
        # We need to run this where the puppeteer config is visible if mermaid-filter looks for it in CWD
        subprocess.run(cmd, check=True, cwd=os.getcwd(), env=env)
        _enforce_table_borders(output_path)
    except subprocess.CalledProcessError as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")
    finally:
        # Cleanup input files
        if os.path.exists(input_path):
            os.remove(input_path)
            if input_path != Path(processed_path) and os.path.exists(processed_path):
                 os.remove(processed_path)

    return FileResponse(
        path=output_path, 
        filename=output_filename, 
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@router.get("/health")
def health_check():
    return {"status": "ok"}
